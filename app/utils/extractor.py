import re
from typing import List
from pathlib import Path

from app.utils.logger import (
    get_logger,
    log_extraction_start,
    log_extraction_success,
)

logger = get_logger("app.utils.extractor")


class DocumentExtractor:
    """Extract text from PDF and DOCX documents"""

    def __init__(self):
        self.supported_extensions = [".pdf", ".docx"]

    def is_supported_file(self, filename: str) -> bool:
        """Check if file extension is supported"""
        return any(filename.lower().endswith(ext) for ext in self.supported_extensions)

    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a document file

        Args:
            file_path (str): Path to the document file

        Returns:
            str: Extracted text

        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file does not exist
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not self.is_supported_file(file_path.name):
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

        logger.info(f"Extracting text from {file_path}")

        if file_path.suffix.lower() == ".pdf":
            return self._extract_pdf_text(file_path)
        elif file_path.suffix.lower() == ".docx":
            return self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

    def _extract_pdf_text(self, file_path: Path) -> str:
        """
        Extract text from PDF using PyMuPDF

        Args:
            file_path (Path): Path to PDF file

        Returns:
            str: Extracted text
        """
        try:
            import fitz

            logger.info(f"Extracting PDF text from {file_path}")

            with fitz.open(str(file_path)) as doc:
                text_parts = []

                for page_num, page in enumerate(doc.pages(), 1):
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(f"\n--- Page {page_num} ---\n")
                        text_parts.append(text)

                full_text = "".join(text_parts)
                logger.info(f"Extracted {len(full_text)} characters from PDF")

                return full_text

        except ImportError:
            raise ImportError(
                "PyMuPDF (fitz) not installed. Please install with: pip install pymupdf"
            )
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise

    def _extract_docx_text(self, file_path: Path) -> str:
        """
        Extract text from DOCX using python-docx

        Args:
            file_path (Path): Path to DOCX file

        Returns:
            str: Extracted text
        """
        try:
            from docx import Document

            logger.info(f"Extracting DOCX text from {file_path}")

            doc = Document(str(file_path))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            full_text = "\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from DOCX")

            return full_text

        except ImportError:
            raise ImportError(
                "python-docx not installed. Please install with: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise


class TextCleaner:
    """Clean and preprocess extracted text"""

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean extracted text by removing common artifacts and normalizing

        Args:
            text (str): Raw extracted text

        Returns:
            str: Cleaned text
        """
        if not text:
            return ""

        text = re.sub(r"\s+", " ", text)
        text = text.strip()

        text = re.sub(r"\n\s*\n\s*\n", "\n\n", text)
        text = re.sub(r"([0-9]+)\s*([,.])\s*([0-9]+)", r"\1\2\3", text)

        text = re.sub(r"\n\d+\s*\n", "\n", text)
        text = re.sub(r"\n[^\n]{1,100}\s*\|\s*[^\n]{1,100}\n", "\n", text)

        text = text.replace("“", '"').replace("”", '"')
        text = text.replace("‘", "'").replace("’", "'")
        text = text.replace("–", "-").replace("—", "-")

        text = re.sub(r"([.!?])\1+", r"\1", text)

        return text.strip()

    @staticmethod
    def remove_resume_noise(text: str) -> str:
        """
        Remove common resume-specific noise

        Args:
            text (str): Cleaned text

        Returns:
            str: Text with resume noise removed
        """
        if not text:
            return ""

        noise_patterns = [
            r"(?i)references available upon request",
            r"(?i)objective:",
            r"(?i)personal information",
            r"(?i)date of birth",
            r"(?i)marital status",
            r"(?i)nationality",
            r"(?i)passport number",
            r"(?i)driver.s license",
        ]

        for pattern in noise_patterns:
            text = re.sub(pattern, "", text, flags=re.IGNORECASE)

        text = re.sub(r"\n([^\n]*@[^\n]*[^\s])\s*\1", r"\n\1", text)
        text = re.sub(r"\n([^\n]*\d{3}[-.\s]?\d{3}[-.\s]?\d{4})\s*\1", r"\n\1", text)

        return text.strip()


class TextChunker:
    """Split text into chunks for processing"""

    def __init__(self, chunk_size: int = 3000, chunk_overlap: int = 300):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks

        Args:
            text (str): Text to chunk
            chunk_size (int): Maximum size of each chunk
            chunk_overlap (int): Overlap between chunks

        Returns:
            List[str]: List of text chunks
        """
        if not text:
            return []

        sentences = re.split(r"[.!?]+", text)
        sentences = [s.strip() for s in sentences if s.strip()]

        chunks = []
        current_chunk = ""

        for sentence in sentences:
            sentence_with_punct = sentence + "."

            if len(current_chunk + sentence_with_punct) <= self.chunk_size:
                current_chunk += " " + sentence_with_punct
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())

                if self.chunk_overlap > 0:

                    overlap_text = current_chunk[-self.chunk_overlap :]
                    current_chunk = overlap_text + " " + sentence_with_punct
                else:
                    current_chunk = sentence_with_punct

        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        logger.info(
            f"Text chunked into {len(chunks)} chunks (size: {self.chunk_size}, overlap: {self.chunk_overlap})"
        )

        return chunks


document_extractor = DocumentExtractor()
text_cleaner = TextCleaner()
text_chunker = TextChunker()


def extract_and_clean_text(file_path: str) -> str:
    """
    Extract and clean text from a document file

    Args:
        file_path (str): Path to the document file

    Returns:
        str: Cleaned extracted text
    """
    log_extraction_start(file_path)

    raw_text = document_extractor.extract_text(file_path)

    cleaned_text = text_cleaner.clean_text(raw_text)
    cleaned_text = text_cleaner.remove_resume_noise(cleaned_text)

    log_extraction_success(len(cleaned_text))

    return cleaned_text


def chunk_for_llm(text: str, chunk_size: int = 3000, overlap: int = 300) -> List[str]:
    """
    Chunk text for LLM processing

    Args:
        text (str): Text to chunk
        chunk_size (int): Size of each chunk
        overlap (int): Overlap between chunks

    Returns:
        List[str]: List of text chunks
    """
    chunker = TextChunker(chunk_size, overlap)
    return chunker.chunk_text(text)
